import os
import uuid
import re
import json
import logging
import pythoncom
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from flask import Flask, request, render_template, redirect, url_for, flash, session, send_file
from werkzeug.utils import secure_filename
from docx2pdf import convert

app = Flask(__name__)
app.secret_key = os.environ.get('FLASK_SECRET_KEY', 'your-secret-key')
#########################
import sys
try:
    import pythoncom
except ImportError:
    pythoncom = None

from docx2pdf import convert

def convert_to_pdf(docx_path, pdf_path):
    if pythoncom and sys.platform == "win32":
        pythoncom.CoInitialize()
        convert(docx_path, pdf_path)
        pythoncom.CoUninitialize()
    else:
        # Fallback (e.g., skip conversion or use alternative)
        print("PDF conversion not supported on this platform")
        # Optionally, implement pdfkit or libreoffice here
##################################################################################################################3
# Directories
app.config['UPLOAD_DIR'] = 'uploads'
app.config['SIGNATURE_DIR'] = os.path.join(app.config['UPLOAD_DIR'], 'signatures').replace('\\', '/')
app.config['DOCX_DIR'] = os.path.join(app.config['UPLOAD_DIR'], 'docx').replace('\\', '/')
app.config['OUTPUT_DIR'] = os.path.join(app.config['UPLOAD_DIR'], 'pdf').replace('\\', '/')
app.config['METADATA_DIR'] = os.path.join(app.config['UPLOAD_DIR'], 'metadata').replace('\\', '/')
app.config['EDIT_HISTORY_DIR'] = os.path.join(app.config['UPLOAD_DIR'], 'edit_history').replace('\\', '/')
app.config['JSON_FILE'] = os.path.join(app.config['UPLOAD_DIR'], 'form_data.json').replace('\\', '/')

# Create directories if they don't exist
for directory in [app.config['UPLOAD_DIR'], app.config['SIGNATURE_DIR'], app.config['DOCX_DIR'], 
                 app.config['OUTPUT_DIR'], app.config['METADATA_DIR'], app.config['EDIT_HISTORY_DIR']]:
    os.makedirs(directory, exist_ok=True)

# Logging setup
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(levelname)s - %(message)s',
                   handlers=[
                       logging.FileHandler(os.path.join(app.config['UPLOAD_DIR'], 'app.log').replace('\\', '/')),
                       logging.StreamHandler()
                   ])
logger = logging.getLogger(__name__)

# Allowed file extensions for signatures
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg'}

# Mock user database
users = {
    'Harsha': '1'
}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def validate_date(date_str):
    try:
        datetime.strptime(date_str, '%Y-%m-%d')
        return True
    except ValueError:
        return False

def sanitize_input(input_str):
    if not input_str:
        return ""
    # Remove potentially harmful characters
    return re.sub(r'[<>]', '', input_str)

def save_signature(data_url, prefix):
    try:
        import base64
        from io import BytesIO
        # Remove data URL prefix
        data = data_url.split(',')[1]
        img_data = base64.b64decode(data)
        filename = f"{prefix}_{uuid.uuid4()}.png"
        filepath = os.path.join(app.config['SIGNATURE_DIR'], filename).replace('\\', '/')
        with open(filepath, 'wb') as f:
            f.write(img_data)
        logger.info(f"Saved signature: {filepath}")
        return filepath
    except Exception as e:
        logger.error(f"Error saving signature: {e}")
        return None

def add_page_number(section):
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    run._r.append(fld)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

def save_edit_history(filename, username, changes):
    history_file = os.path.join(app.config['EDIT_HISTORY_DIR'], f"{filename}_history.json").replace('\\', '/')
    history = []
    if os.path.exists(history_file):
        try:
            with open(history_file, 'r') as f:
                history = json.load(f)
        except Exception as e:
            logger.error(f"Error loading edit history: {e}")
    
    history.append({
        'timestamp': datetime.now().isoformat(),
        'username': username,
        'changes': changes
    })
    
    try:
        with open(history_file, 'w') as f:
            json.dump(history, f, indent=4)
        logger.info(f"Saved edit history for {filename}")
    except Exception as e:
        logger.error(f"Error saving edit history: {e}")

def create_document(content):
    """Generate the Chervic Master Agreement document."""
    try:
        doc = Document()
        
        # Set document margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Load edited lines if available
        edited_lines = {}
        filename = content.get('pdf_filename', '')
        if filename:
            lines_file = os.path.join(app.config['EDIT_HISTORY_DIR'], f"{filename}_lines.json").replace('\\', '/')
            if os.path.exists(lines_file):
                try:
                    with open(lines_file, 'r') as f:
                        edited_lines = json.load(f)
                except Exception as e:
                    logger.error(f"Error loading line edits: {e}")
        
        line_id = 0
        def get_line_text(original_text):
            nonlocal line_id
            lines = original_text.split('\n')
            new_lines = []
            for line in lines:
                if line.strip():
                    line_key = f'line_{line_id}'
                    new_lines.append(edited_lines.get(line_key, line.strip()))
                    line_id += 1
                else:
                    new_lines.append(line)
            return '\n'.join(new_lines)
        
        content["currency"] = "☑ " + content.get('currency', 'USD') + " ☐ " + ("USD" if content.get('currency', 'USD') == "INR" else "INR")
        
        # Helper functions for document formatting
        def add_heading(text, level):
            heading = doc.add_heading(get_line_text(text), level=level)
            run = heading.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            return heading
        
        def add_paragraph(text):
            para = doc.add_paragraph(get_line_text(text))
            para.style.font.name = 'Times New Roman'
            para.style.font.size = Pt(12)
            return para
        
        def add_table(rows, cols, headers, data=None, style='Table Grid'):
            logger.info(f"Creating table with {rows} rows, {cols} cols, headers: {headers}")
            table = doc.add_table(rows=rows, cols=cols)
            table.style = style
            for i, header in enumerate(headers):
                if i < len(table.rows[0].cells):
                    cell = table.rows[0].cells[i]
                    cell.text = get_line_text(header)
                    cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                    cell.paragraphs[0].runs[0].font.size = Pt(12)
                else:
                    logger.warning(f"Header index {i} exceeds table columns ({cols})")
            if data:
                for i, row_data in enumerate(data, 1):
                    row_data = list(row_data) + [''] * (cols - len(row_data)) if len(row_data) < cols else row_data[:cols]
                    logger.debug(f"Populating table row {i}: {row_data}")
                    for j, cell_data in enumerate(row_data):
                        if i < len(table.rows) and j < len(table.rows[i].cells):
                            cell = table.rows[i].cells[j]
                            cell.text = get_line_text(str(cell_data))
                            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                            cell.paragraphs[0].runs[0].font.size = Pt(12)
                        else:
                            logger.warning(f"Invalid cell access at row {i}, col {j}")
            return table

        # Cover Page
        para = doc.add_paragraph(get_line_text("Chervic Advisory Services – Master Agreement"))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        
        para = doc.add_paragraph(get_line_text("For ADA – Ariba Digital Assistant & AGS – Ariba Global Support"))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        para = doc.add_paragraph(get_line_text(f"Company: {content.get('customer_name', '')}\nDate: {content.get('agreement_date', '')}\nVersion: 1.0"))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.italic = True
        
        doc.add_section(WD_SECTION.NEW_PAGE)
        add_page_number(doc.sections[-1])
        
        # Master Agreement
        add_heading("Master Agreement", 1)
        
        add_heading("1. Parties & Purpose", 2)
        add_paragraph(f"This Agreement (“Agreement”) is made between Chervic Advisory Services (“Chervic”) and {content.get('customer_name', '')} (“Customer”) to license and support ADA and/or AGS.\n*Company details automatically downloaded from the Company registration details.*")
        
        add_heading("2. Plans & Fees", 2)
        add_paragraph("*Automatically downloaded from subscription plan*")
        table = add_table(4, 5, ["Plan", "Scope", "Billing", "Change Rate**", "Included Support*"], [
            ["PayAsYouUse", "", "", "", ""],
            ["Subscription", "", "", "", ""],
            ["Support", "", "", "", ""]
        ])
        add_paragraph("*Support tiers and SLAs in Exhibit B & C. Full pricing appears in the Schedule (Exhibit A).\n**Change rate suggests rate charged per hour if a ticket takes more than 8 hours to resolve.")
        
        add_heading("3. Key Licence Terms", 2)
        add_paragraph("- Nonexclusive, nontransferable CAS licence for internal use only.\n- No title or IP is transferred; Customer receives user rights only for the Subscription Term.\n- Competitors may not benchmark, reverse-engineer, or resell services.")
        
        add_heading("4. Customer Obligations", 2)
        add_paragraph("1. **Designated Contacts**: Up to three authorized users may raise support cases.\n2. **Assistance**: Timely access to systems, test data, and subject-matter experts.\n3. **Compliance**: Observe export, data protection, and anti-corruption laws.\n4. **Security**: Keep credentials safe; report breaches within 24 hours.\n5. **Content Warranty**: Customer data must be lawful and non-infringing.")
        
        add_heading("5. Background Verification (BGV) & Database (DB) Access", 2)
        add_paragraph("- **Chervic Personnel**: All project staff undergo 3-year employment, education, criminal record, and ID checks. Results retained for audit.\n- **DB Perspective**:\n  - Production DB access is read-only & least-privilege; credentials rotate every 90 days.\n  - Sensitive data at rest is AES-256 encrypted; in transit via TLS 1.3.\n  - Access logs kept 365 days and shared with Customer on request.")
        
        add_heading("6. Non-Solicitation (5 Years)", 2)
        add_paragraph("Customer shall not hire, contract, or solicit any Chervic employee/contractor engaged on this project during the term and for five (5) years after final service delivery.")
        
        add_heading("7. Service Levels", 2)
        add_paragraph("Refer to Exhibit C.")
        
        add_heading("8. Warranties & Liability", 2)
        add_paragraph("- Chervic warrants services will perform materially per documentation and industry practice.\n- **Cap**: Aggregate liability limited to fees paid in the preceding 12 months.\n- No liability for indirect or consequential damages.")
        
        add_heading("9. Term, Renewal & Exit", 2)
        add_paragraph("- **Initial Subscription Term**: 12 months (unless stated otherwise).\n- Auto-renews for like periods unless either party gives 30-day notice.\n- On termination, Customer may export data for 30 days; Chervic deletes remaining data 60 days post-termination (except one encrypted archive for legal purposes).")
        
        add_heading("10. Payment & Price Changes", 2)
        add_paragraph("- Fees non-cancellable/non-refundable once invoiced.\n- **Annual price uplift**: ≤ 7% unless otherwise agreed 30 days before renewal.\n- Late payments accrue 1.5% per month.")
        
        add_heading("11. Confidentiality", 2)
        add_paragraph("- Each party protects the other’s Confidential Information with at least the same care it uses for its own, for five (5) years (software & source code – perpetual).\n- **Permitted disclosures**: Legal advisors, auditors, mandatory law.")
        
        add_heading("12. Governing Law & Dispute Resolution", 2)
        add_paragraph("- **Governing Law**: India – Courts of New Delhi.\n- Escalate disputes to VP-level within 15 days; if unresolved, refer to arbitration under Indian Arbitration & Conciliation Act, 1996.")
        
        add_heading("Signature Page", 2)
        add_paragraph(f"**Chervic Advisory Services**\nName: {content.get('chervic_name', '_______________________')}\nTitle: {content.get('chervic_title', '_______________________')}\nDate: {content.get('chervic_date', '_______________________')}")
        if content.get('chervic_signature') and os.path.exists(content['chervic_signature']):
            try:
                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(content['chervic_signature'], width=Inches(2))
                logger.info(f"Embedded Chervic signature: {content['chervic_signature']}")
            except Exception as e:
                logger.error(f"Error embedding Chervic signature: {e}")
                add_paragraph("Chervic Signature: [Error loading signature]")
        
        add_paragraph(f"\n**Customer**\nName: {content.get('customer_sign_name', '_______________________')}\nTitle: {content.get('customer_sign_title', '_______________________')}\nDate: {content.get('customer_sign_date', '_______________________')}")
        if content.get('customer_signature') and os.path.exists(content['customer_signature']):
            try:
                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(content['customer_signature'], width=Inches(2))
                logger.info(f"Embedded Customer signature: {content['customer_signature']}")
            except Exception as e:
                logger.error(f"Error embedding Customer signature: {e}")
                add_paragraph("Customer Signature: [Error loading signature]")
        
        add_heading("Attachments", 2)
        add_paragraph("- Exhibit A – Plan & Pricing Schedule\n- Exhibit B – Support & Maintenance Matrix\n- Exhibit C – SLA & Service Credit Table")
        
        # Exhibit A
        doc.add_section(WD_SECTION.NEW_PAGE)
        add_page_number(doc.sections[-1])
        
        add_heading("Exhibit A: Plan & Pricing Schedule", 1)
        add_paragraph(f"Chervic Advisory Services | ADA & AGS CAS Solutions\nThis Schedule A-___ forms an integral part of the Master CAS Agreement dated {content.get('agreement_date', '')} between Chervic Advisory Services (“Chervic”) and {content.get('customer_name', '')} (“Customer”).")
        
        add_heading("1. Selected Plan", 2)
        add_paragraph("Please tick one or more of the options below:")
        table = add_table(4, 2, ["Plan Type", "Selected?"], [
            ["Pay-As-You-Use", "☐"],
            ["Subscription", "☐"],
            ["Support-Only Plan", "☐"]
        ])
        
        add_heading("2. Plan Components & Pricing", 2)
        add_heading("A. Pay-As-You-Use Pricing", 3)
        table = add_table(7, 3, ["Service Item", "Unit", "Rate (USD)"], [
            ["ADA Bot Interaction", "Per query", "$0.12 / interaction"],
            ["Complaint Ticket Resolution", "Per ticket", "$1.50 / ticket"],
            ["Change Rate*", "", "$150 / hour"],
            ["BOT X", "Per run", "$0.50 / workflow"],
            ["BOT Y", "Per call", "$0.08 / call"],
            ["BOT Z", "Per file", "$0.30 / file"]
        ])
        add_paragraph("*Change rate suggests rate charged per hour if a ticket takes more than 8 hours to resolve.\nBilled monthly, minimum invoice value: $100/month.")
        
        add_heading("B. Annual Subscription Packages", 3)
        table = add_table(7, 3, ["Service Item", "Unit", "Rate (USD)"], [
            ["ADA Bot Interaction", "Per query", "$0.12 / interaction"],
            ["Complaint Ticket Resolution", "Per ticket", "$1.50 / ticket"],
            ["Change Rate*", "", "$150 / hour"],
            ["BOT X", "Per run", "$0.50 / workflow"],
            ["BOT Y", "Per call", "$0.08 / call"],
            ["BOT Z", "Per file", "$0.30 / file"]
        ])
        add_paragraph("All plans include Basic Support (see Exhibit B). Additional support levels can be added separately.")
        
        add_heading("C. Support-Only Plans", 3)
        table = add_table(7, 3, ["Service Item", "Unit", "Rate (USD)"], [
            ["ADA Bot Interaction", "Per query", "$0.12 / interaction"],
            ["Complaint Ticket Resolution", "Per ticket", "$1.50 / ticket"],
            ["Change Rate*", "", "$150 / hour"],
            ["BOT X", "Per run", "$0.50 / workflow"],
            ["BOT Y", "Per call", "$0.08 / call"],
            ["BOT Z", "Per file", "$0.30 / file"]
        ])
        
        add_heading("3. Term and Renewal", 2)
        add_paragraph(f"- **Start Date**: {content.get('start_date', '')}\n- **End Date**: {content.get('end_date', '')}\n- Auto-renews for additional 12-month terms unless terminated with 30 days’ written notice.")
        
        add_heading("4. Billing Information", 2)
        add_paragraph("*Billing details automatically downloaded from the Company registration details.*")
        table = add_table(7, 2, ["Field", "Entry"], [
            ["Billing Contact Name", content.get('billing_contact_name', '')],
            ["Billing Email", content.get('billing_email', '')],
            ["Department", content.get('department', '')],
            ["Billing Address", content.get('billing_address', '')],
            ["GSTIN / Tax ID", content.get('tax_id', '')],
            ["PO Number (if applicable)", content.get('po_number', '')],
            ["Preferred Currency", content.get('currency', '')]
        ])
        
        add_heading("5. Signatures", 2)
        add_paragraph(f"**Chervic Advisory Services**\nName: {content.get('chervic_name', '_______________________')}\nTitle: {content.get('chervic_title', '_______________________')}\nDate: {content.get('chervic_date', '_______________________')}")
        if content.get('chervic_signature') and os.path.exists(content['chervic_signature']):
            try:
                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(content['chervic_signature'], width=Inches(2))
                logger.info(f"Embedded Chervic signature in Exhibit A: {content['chervic_signature']}")
            except Exception as e:
                logger.error(f"Error embedding Chervic signature in Exhibit A: {e}")
                add_paragraph("Chervic Signature: [Error loading signature]")
        
        add_paragraph(f"\n**Customer**\nName: {content.get('customer_sign_name', '_______________________')}\nTitle: {content.get('customer_sign_title', '_______________________')}\nDate: {content.get('customer_sign_date', '_______________________')}")
        if content.get('customer_signature') and os.path.exists(content['customer_signature']):
            try:
                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(content['customer_signature'], width=Inches(2))
                logger.info(f"Embedded Customer signature in Exhibit A: {content['customer_signature']}")
            except Exception as e:
                logger.error(f"Error embedding Customer signature in Exhibit A: {e}")
                add_paragraph("Customer Signature: [Error loading signature]")
        
        # Exhibit B
        doc.add_section(WD_SECTION.NEW_PAGE)
        add_page_number(doc.sections[-1])
        
        add_heading("Exhibit B: Support & Maintenance Services", 1)
        add_paragraph("Chervic Advisory Services\nThis exhibit outlines the scope of Support & Maintenance Services provided under the Master CAS Agreement and associated subscription plan (see Exhibit A).")
        
        add_heading("1. Included Support Services", 2)
        add_paragraph("All Customers are entitled to the following services, subject to their selected Support Tier:")
        table = add_table(7, 3, ["Support Feature", "", ""], [
            ["24×7 technical assistance", "", ""],
            ["Response time SLAs", "", ""],
            ["Dedicated Technical Account Manager (TAM)", "", ""],
            ["Quarterly service reviews", "", ""],
            ["Root cause analysis documentation", "", ""],
            ["Access to self-service portal", "", ""]
        ])
        add_paragraph("**Business Hours**: Monday–Friday, 9:30 AM to 6:30 PM IST (GMT+5:30), excluding national holidays.")
        
        add_heading("2. Support Portal Access", 2)
        add_paragraph("Chervic provides 24×7 access to its online Support Portal at:\n🔗 https://support.chervic.com (customizable domain)\n**Features**:\n- Submit & manage support tickets\n- View status & history of requests\n- Access product documentation, FAQs, release notes")
        
        add_heading("3. Designated Customer Contacts", 2)
        add_paragraph("The Customer may designate authorized personnel who can log support tickets and receive updates. Changes must be submitted in writing to Chervic’s support desk.")
        
        add_heading("4. Exclusions", 2)
        add_paragraph("Chervic shall not be responsible for:\n- Issues caused by unauthorized changes or use beyond documented scope\n- Failures in Customer’s hardware, local software, or internet connectivity\n- Third-party software not supplied or integrated by Chervic\n- Problems resulting from modifications made without prior approval")
        
        # Exhibit C
        doc.add_section(WD_SECTION.NEW_PAGE)
        add_page_number(doc.sections[-1])
        
        add_heading("Exhibit C: Service Level Agreement (SLA)", 1)
        add_paragraph("Chervic Advisory Services – ADA & AGS CAS Services\nThis Exhibit defines the service availability commitments, credit entitlements, and performance exclusions applicable to CAS Services under the Master Agreement.")
        
        add_heading("1. System Availability Commitment", 2)
        add_paragraph("Chervic commits to a System Availability of ≥ 99.0% annually for the production environment of ADA and AGS services.\n**Definition**:\nSystem Availability = Total uptime minutes during the year ÷ (Total minutes in the year – Excluded downtime) × 100")
        
        add_heading("2. Exclusions from Availability Calculation", 2)
        add_paragraph("Downtime shall not be counted against SLA if caused by:\n- Scheduled Maintenance, with at least 48 hours’ notice\n- Force Majeure events (e.g., natural disasters, cyberattacks, war, etc.)\n- Malicious activity not attributable to Chervic’s infrastructure\n- Customer-side failures, including:\n  - Local area network or internet connectivity\n  - Device/browser misconfigurations\n  - Delays due to user mismanagement\n  - Third-party software/services not managed by Chervic\n  - Customer-imposed constraints, access restrictions, or security blockages")
        
        add_heading("3. Scheduled Maintenance", 2)
        add_paragraph("- Performed during off-peak hours (typically Saturday/Sunday, 10:00 PM–6:00 AM IST)\n- Maximum of 8 hours per month (unless urgent)\n- Notified via portal or email at least 2 business days in advance")
        
        add_heading("4. SLA Credits", 2)
        add_paragraph("If Chervic fails to meet the annual availability target of 99.0%, the Customer is eligible for service credits:")
        table = add_table(5, 2, ["Actual Availability", "Credit (as % of Annual Fee for Affected Service)"], [
            ["98.0% – 98.99%", "5%"],
            ["95.0% – 97.99%", "10%"],
            ["90.0% – 94.99%", "20%"],
            ["Below 90%", "30% (max)"]
        ])
        add_paragraph("- **Maximum Credit**: Not to exceed 30% of the annual fee\n- **Credit Application**: Applied to the next invoice or term extension\n- **Claim Window**: Customer must submit a written credit request within 30 days of year-end")
        
        add_heading("5. Sole Remedy", 2)
        add_paragraph("Service credits defined in this SLA are the Customer’s sole and exclusive remedy for Chervic’s failure to meet system availability targets.")
        
        add_heading("6. Monitoring & Reporting", 2)
        add_paragraph("- System uptime and incident logs are continuously monitored\n- Monthly availability summaries available upon request\n- Real-time issue notifications provided via support portal or email")
        
        # CAS Agreement
        doc.add_section(WD_SECTION.NEW_PAGE)
        add_page_number(doc.sections[-1])
        
        add_heading("CAS Agreement", 1)
        
        add_heading("PLEASE READ THIS AGREEMENT CAREFULLY", 2)
        add_paragraph("*BY ACCESSING OR USING CHERVIC ADVISORY SERVICES’ (“CHERVIC”) SOFTWARE AND SERVICES, INCLUDING THE ARIBA DIGITAL ASSISTANT (ADA) AND ARIBA GLOBAL SUPPORT (AGS) SERVICE OFFERINGS, YOU (THE “CUSTOMER”) ACKNOWLEDGE AND AGREE TO BE LEGALLY BOUND BY THE TERMS AND CONDITIONS OF THIS AGREEMENT. IF YOU DO NOT AGREE TO THESE TERMS, YOU MUST NOT ACCESS OR USE CHERVIC’S SERVICES.*\n*IF YOU HAVE A FULLY EXECUTED MASTER AGREEMENT OR STATEMENT OF WORK (SOW) THAT EXPRESSLY GOVERNS YOUR ACCESS TO ADA OR AGS, THAT DOCUMENT SHALL TAKE PRECEDENCE OVER THIS AGREEMENT.*")
        
        add_heading("1. DEFINITIONS", 2)
        add_paragraph("- **Administrator User**: Each employee or representative of the Customer designated to administer and manage the ADA or AGS services. These individuals may require onboarding, enablement, or training as prescribed by Chervic.\n- **Customer Content**: All information, documentation, credentials, business process inputs, and materials provided by the Customer to Chervic for service delivery, including Ariba realm access, support tickets, invoice formats, approval flows, and custom integration documentation.\n- **Documentation**: User manuals, product specifications, training guides, online help materials, and related documentation provided by Chervic for ADA and AGS services.\n- **Host**: Cloud-based or on-prem infrastructure from which ADA and AGS services are delivered, owned, managed, or subcontracted by Chervic or its affiliates.\n- **Maintenance Services**: Continuous monitoring, updates, fixes, enhancements, and ongoing functional support provided by Chervic to ensure stable operation.\n- **Other Services**: Custom implementation, system integration, advisory, training, or consulting services provided by Chervic outside the standard scope of ADA and AGS, billed separately under time and material or fixed-bid SOWs.\n- **Schedule or Exhibit**: Attached or referenced document outlining the scope of work, pricing, subscription details, or additional services.\n- **Software**: Object code and application logic powering ADA and AGS, including bot frameworks, dashboards, integration tools, and updates or releases deployed by Chervic.\n- **Statement of Work (SOW)**: Mutually executed document defining specific deliverables, timelines, and objectives for professional or customized services.\n- **CAS Services**: Chervic’s cloud-hosted automation and support platforms, specifically ADA (Ariba Digital Assistant) and AGS (Ariba Global Support), accessed remotely and licensed on a term-use basis.\n- **Subscription Term**: Duration specified in the Schedule during which the Customer is authorized to access and use ADA and/or AGS. Auto-renews for additional 12-month periods unless either party gives 30 days’ prior written notice of non-renewal.")
        
        add_heading("2. CAS SERVICES", 2)
        add_heading("2.1", 3)
        add_paragraph(f"During the Subscription Term, the Customer is granted a non-exclusive, non-transferable, royalty-free, limited worldwide license to access and use the Chervic ADA and/or AGS services solely for internal business operations. Usage is subject to the terms and conditions of this Agreement and restricted to the number of service units, users, or features defined in the applicable Schedule or Statement of Work.")
        
        add_heading("2.2", 3)
        add_paragraph("The Customer acknowledges that this Agreement provides only a subscription to use the service, and no ownership or additional rights are transferred other than those explicitly described in Section 2.1. The underlying software, bots, frameworks, integrations, and infrastructure remain the sole property of Chervic Advisory Services.")
        
        add_heading("3. FREE TRIAL", 2)
        add_paragraph("If the Customer registers for a Free Trial of ADA and/or AGS via Chervic’s website or pre-authorized portal, Chervic may provide access to selected CAS services on a trial basis free of charge until the earliest of:\n- (a) The end of the applicable trial period;\n- (b) The effective start date of a paid Subscription Term; or\n- (c) Termination of the trial at Chervic’s sole discretion.\nAdditional trial-specific conditions may appear on the registration page or in email communications, incorporated by reference into this Agreement.")
        
        add_heading("Important Trial Limitations", 3)
        add_paragraph("- **Data Loss**: Any data uploaded or generated, and customizations performed during the Free Trial, will be permanently lost unless the Customer subscribes to a paid version equal to or greater than the trial configuration.\n- **No Downgrade Migration**: Data or settings cannot be migrated to a lower-tier plan post-trial. Customers must export data before the trial ends.\n- **No Warranties**: Services during the Free Trial are provided “as-is”, without warranties, service-level commitments, or indemnification.\n- **Limited Liability**: Chervic’s liability during the Free Trial is capped at INR ₹8,000 / USD $100, or the minimum enforceable amount under applicable law.\n- **Customer Responsibility**: The Customer is fully responsible for any loss, breach, or misuse of the CAS Services during the trial, including full indemnification for damages caused to Chervic or third parties.")
        
        add_heading("Trial Restrictions", 3)
        add_paragraph("- Competitors of Chervic are prohibited from accessing any Free Trial without prior written consent.\n- Trials must not be used for benchmarking, monitoring, reverse engineering, or competitive analysis.")
        
        add_heading("4. RESTRICTIONS", 2)
        add_paragraph("The Customer agrees not to, and shall not permit any third party to:\ni. Copy, reproduce, or republish the ADA or AGS services, including supporting Software;\nii. Share or provide access to the CAS Services to any individual or entity other than authorized users;\niii. Use the CAS Services to offer time-sharing, hosting, managed services, or service bureau functions to third parties;\niv. Modify, adapt, or create derivative works based on ADA, AGS, or their supporting documentation;\nv. Remove, alter, or obscure proprietary notices, copyright markings, or branding;\nvi. Reverse engineer, decompile, disassemble, or attempt to derive source code from Chervic software components;\nvii. Access or use the CAS Services or Documentation to develop or enhance a competitive or similar service or software.\nAll rights, titles, and interests in ADA and AGS software, documentation, services, enhancements, improvements, feedback, and intellectual property remain solely with Chervic. Any suggestions, custom developments, or modifications made during the Subscription Term may be retained and used by Chervic, with Customer assigning any rights in such contributions without further consideration.")
        
        add_heading("5. CUSTOMER RESPONSIBILITIES", 2)
        add_heading("5.1 Assistance", 3)
        add_paragraph("The Customer shall provide timely access to personnel, infrastructure, credentials, documentation, and other reasonable cooperation to enable Chervic to deliver services efficiently. The success of ADA and AGS implementation depends on the accuracy and timeliness of this support.")
        
        add_heading("5.2 Compliance with Laws", 3)
        add_paragraph("The Customer shall ensure that its use of ADA and AGS complies with all applicable national and international laws, including data privacy, electronic communications, and cross-border data transfers. The Customer agrees not to upload or transmit content that infringes third-party intellectual property rights or violates privacy or data protection laws.")
        
        add_heading("5.3 Unauthorized Use & False Identity", 3)
        add_paragraph("The Customer shall:\n- (a) Notify Chervic immediately of unauthorized access or suspected credential breaches;\n- (b) Take reasonable efforts to halt unauthorized usage;\n- (c) Not misrepresent identity or provide false credentials to access ADA or AGS.")
        
        add_heading("5.4 Administrator User Responsibility", 3)
        add_paragraph("The Customer assumes full responsibility for the conduct, accuracy, and actions of its designated Administrator Users. Chervic is not liable for damage or data loss resulting from their mismanagement or negligence.")
        
        add_heading("5.5 Customer Content Input", 3)
        add_paragraph("The Customer is solely responsible for all data, documents, and content entered into ADA or AGS, ensuring it:\n- Does not violate intellectual property rights;\n- Is not defamatory, obscene, offensive, or malicious;\n- Complies with all applicable laws and policies.\nBreaches may lead to termination of access without refund.")
        
        add_heading("5.6 License to Customer Content", 3)
        add_paragraph("The Customer grants Chervic a limited, non-exclusive, non-transferable license to access, host, copy, display, process, and transmit Customer Content solely for delivering and maintaining ADA and AGS.")
        
        add_heading("5.7 Ownership and IP Rights", 3)
        add_paragraph("The Customer retains ownership of its proprietary data, business logic, and uploaded content. Chervic retains ownership of all tools, services, configurations, and intellectual property related to ADA and AGS. Third-party technologies (e.g., Microsoft Teams, SAP connectors) are governed by their respective license terms.")
        
        add_heading("5.8 Suggestions and Feedback", 3)
        add_paragraph("Feedback, ideas, enhancement requests, or suggestions provided by the Customer may be used by Chervic without restriction. Chervic retains a worldwide, perpetual, royalty-free license to use and incorporate such feedback without obligation.")
        
        add_heading("6. ORDERS AND PAYMENT", 2)
        add_heading("6.1 Orders", 3)
        add_paragraph("The Customer shall acquire access to ADA and/or AGS through an executed Schedule or Statement of Work (SOW). All services are governed solely by this CAS Agreement and the relevant Schedule. In case of conflict, the Schedule prevails.")
        
        add_heading("6.2 Invoicing and Payment", 3)
        add_paragraph("Unless specified in the Schedule:\n- Chervic will invoice on the effective date of the Schedule or service commencement.\n- Invoices are payable within 30 days.\n- Payments in Indian Rupees (INR) or U.S. Dollars (USD) as specified.\n- Overdue payments accrue 1.5% interest per month.\n**Additional Conditions**:\n- Fees are based on subscribed services, not actual usage.\n- Payment obligations are non-cancellable, and fees are non-refundable.\n- Service unit quantities cannot be reduced during the active term.\n- Per-unit pricing may increase up to 7% annually, unless communicated 30 days prior to renewal.\n- Promotional pricing applies only for the initial term, reverting to standard pricing upon renewal.\nIf the Customer reduces scope or volume in a renewal term, Chervic may recalculate unit pricing.")
        
        add_heading("6.3 Expenses", 3)
        add_paragraph("For professional or onsite services not covered in the standard plan, the Customer shall reimburse Chervic for reasonable travel and accommodation expenses, subject to prior written approval and adherence to Customer’s documented travel policies.")
        
        add_heading("6.4 Taxes", 3)
        add_paragraph("Prices exclude applicable taxes. Chervic will include GST, VAT, sales tax, or similar charges on invoices. The Customer is responsible for all taxes related to purchase or use, excluding Chervic’s net income or corporate tax obligations.")
        
        add_heading("7. TERM AND TERMINATION", 2)
        add_heading("7.1 Term", 3)
        add_paragraph("This Agreement is effective on the Effective Date and remains in force until the expiry or termination of the Subscription Term. Unless stated in the Schedule, the term auto-renews for successive periods equal to the previous term or 12 months, unless either party provides 30 days’ prior written notice of non-renewal.")
        
        add_heading("7.2 Termination for Cause", 3)
        add_paragraph("Either party may terminate with 30 days’ written notice if the other commits a material breach and fails to remedy it within that period. For Chervic, material breach includes non-compliance with SOW, Schedule, or Agreement obligations.")
        
        add_heading("7.3 Suspension for Non-Payment", 3)
        add_paragraph("Chervic may suspend services if undisputed dues remain unpaid 15 days after written notice. Suspension does not absolve payment obligations, and Chervic is not liable for losses due to suspension.")
        
        add_heading("7.4 Suspension Due to Harmful Use", 3)
        add_paragraph("Chervic may suspend services if the Customer’s use causes immediate or ongoing harm to Chervic systems or third parties. Chervic will notify the Customer and work to resolve the issue, with no liability for losses.")
        
        add_heading("7.5 Effect of Termination", 3)
        add_paragraph("Upon termination:\n- (a) Chervic discontinues access to CAS Services, and all granted rights terminate.\n- (b) If due to Customer breach, Customer pays all fees due through the Subscription Term.\n- (c) If due to Chervic breach, Chervic refunds pre-paid, unused fees.\n- (d) In other cases, no refunds are issued, and Customer remains liable for committed fees.\nUpon request, the receiving party shall return or destroy confidential information, certifying destruction, though legal counsel may retain one archival copy.")
        
        add_heading("8. SERVICE LEVEL AGREEMENT (SLA)", 2)
        add_paragraph("The SLA for ADA and AGS services is defined in Exhibit C, outlining availability guarantees, support response times, and performance obligations. Service credits are the Customer’s sole remedy for unavailability or performance degradation, unless stated in the Schedule.")
        
        add_heading("9. WARRANTIES", 2)
        add_heading("9.1 Service Warranty", 3)
        add_paragraph("Chervic warrants that ADA and AGS services will be delivered professionally, consistent with industry standards, and will function as described in the Documentation. The sole remedy for breach is termination and applicable refunds per Section 7.")
        
        add_heading("9.2 Service Limitations", 3)
        add_paragraph("Chervic warrants services will function materially as described, but:\n- No guarantee of error-free, uninterrupted service or immunity to bugs/downtime.\n- Chervic is not liable for delays/outages due to internet or third-party infrastructure.\n- Services may experience temporary latency due to cloud infrastructure.\nThis section is the sole warranty. No implied warranties (merchantability, fitness for purpose, non-infringement) apply. Chervic and vendors are not responsible for:\n- Data loss/destruction by third parties or external attacks,\n- Unauthorized access/modification by Customer’s users,\n- Failures from platform misuse or modifications.")
        
        add_heading("10. LIMITATION OF LIABILITY", 2)
        add_paragraph("To the fullest extent permitted by law:\n- Neither party is liable for indirect, incidental, special, or consequential damages (e.g., data loss, revenue, profits, business interruption), even if advised of the possibility.\n- Total liability is capped at fees paid by the Customer in the 12 months preceding the claim.\n- This limitation does not apply to breaches of Confidentiality, Restrictions on Use, or indemnification obligations.")
        
        add_heading("11. INDEMNIFICATION", 2)
        add_heading("11.1 By Chervic Advisory Services", 3)
        add_paragraph("Chervic shall indemnify the Customer against third-party claims alleging that ADA or AGS infringes patents, copyrights, trademarks, or trade secrets, provided services are used per this Agreement and Documentation. Chervic also indemnifies for claims involving bodily injury or death caused by its negligence or willful misconduct.\nChervic has no obligation for claims from:\n- Customer-provided data/content,\n- Unauthorized modifications,\n- Use outside permitted scope.\nChervic may: procure continued access, modify services to resolve infringement, or terminate the Subscription Term with a refund of unused fees.")
        
        add_heading("11.2 By Customer", 3)
        add_paragraph("The Customer shall indemnify Chervic against third-party claims that Customer’s Content or use of services:\n- Infringes IP rights,\n- Violates applicable law,\n- Results from negligence or misuse.")
        
        add_heading("11.3 Indemnification Conditions", 3)
        add_paragraph("- Indemnified party must promptly notify the indemnifying party of claims.\n- Indemnifying party controls defense and settlement.\n- Indemnified party provides reasonable assistance at the indemnifying party’s expense.")
        
        add_heading("12. CONFIDENTIALITY", 2)
        add_heading("12.1 Definition", 3)
        add_paragraph("“Confidential Information” includes non-public information:\n(a) Marked or identified as confidential;\n(b) Verbally disclosed and confirmed in writing within 30 days;\n(c) Reasonably understood as confidential;\n(d) Identified as such under this Agreement;\n(e) Third-party confidential information.\nIncludes trade secrets, business plans, pricing, system architecture, employee/contractor data, BGV records, technical specifications.\n- Customer Content is Customer’s Confidential Information.\n- ADA & AGS platforms, documentation, and source code are Chervic’s Confidential Information.")
        
        add_heading("12.2 Protection Obligations", 3)
        add_paragraph("For the Agreement’s duration and 5 years after (perpetual for software/source code):\n- Maintain confidentiality with at least the same care as own data (no less than reasonable care).\n- Use information only to fulfill Agreement obligations.\n- Not disclose to third parties unless necessary and bound by equivalent confidentiality.\nDisclosure allowed to legal, accounting, or technical advisors under strict confidentiality.\nNo reverse engineering, disassembly, or decompilation of software or prototypes.")
        
        add_heading("12.3 Exclusions", 3)
        add_paragraph("Confidential Information excludes information that:\n- Becomes public without breach;\n- Was lawfully known to the receiving party;\n- Is received from another source without restriction;\n- Is independently developed without using Confidential Information.\nIf legally compelled to disclose, the receiving party must notify the other and cooperate in seeking protective measures.\nParties may disclose the Agreement’s existence and general business relationship description, but not terms or pricing, unless to trusted professionals under confidentiality.")
        
        add_heading("13. GENERAL PROVISIONS", 2)
        add_heading("13.1 Non-Exclusive Service", 3)
        add_paragraph("ADA and AGS are multi-tenant platforms offered non-exclusively. Chervic may offer similar services to other clients.")
        
        add_heading("13.2 Personal Data and Data Protection", 3)
        add_paragraph("Chervic may process, transmit, and store personal data (e.g., BGV, email IDs, credentials) to fulfill Agreement obligations. The Customer, as Data Controller, ensures:\n- Employee consent and disclosures are obtained;\n- No unlawful/restricted data (e.g., biometrics, financials, health records) are transferred without instruction;\n- Compliance with data privacy laws (e.g., India’s DPDP Act, GDPR).\nChervic processes data only as instructed and not for other purposes.")
        
        add_heading("13.3 Chervic's Data Handling Commitment", 3)
        add_paragraph("Chervic will:\n- Adhere to its Data Protection & Security Policy (available on request);\n- Use authorized data processors in compliant jurisdictions;\n- Maintain confidentiality, access controls, encryption, and regular BGV for employees;\n- Notify Customer immediately of data breaches.\nCustomer retains ownership of data and is responsible for its accuracy, legality, and compliance.")
        
        add_heading("13.4 Non-Solicitation (Hiring Clause)", 3)
        add_paragraph("Customer agrees not to solicit, recruit, hire, or engage Chervic personnel involved in service delivery for 5 years from their last engagement. Violation incurs a liquidated damage fee of INR 25 lakhs or USD $30,000 per individual, payable within 30 days.")
        
        add_heading("13.5 Assignment", 3)
        add_paragraph("Neither party may assign rights or delegate obligations without consent, except in mergers, asset sales, or to affiliates/subsidiaries maintaining obligations. Subcontracting is permitted, with the primary party responsible for performance.")
        
        add_heading("13.6 Force Majeure", 3)
        add_paragraph("Neither party is liable for performance failures (excluding payment) due to uncontrollable causes (e.g., natural disasters, cyberattacks, strikes). The affected party must notify promptly and resume performance ASAP.")
        
        add_heading("13.7 Waiver", 3)
        add_paragraph("Waivers must be written and signed. Failure to exercise a right is not a waiver.")
        
        add_heading("13.8 Severability", 3)
        add_paragraph("Invalid provisions shall be reformed to reflect original intent, with the remainder of the Agreement remaining in effect.")
        
        add_heading("13.9 Entire Agreement", 3)
        add_paragraph("This Agreement, with Schedules and Exhibits, is the entire agreement, superseding prior communications. No other terms (e.g., in purchase orders) apply unless accepted in writing.")
        
        add_heading("13.10 Survival", 3)
        add_paragraph("Sections 4, 7, 9, 10, 11, 12, and 13 survive termination.")
        
        add_heading("13.11 Publicity", 3)
        add_paragraph("With approval, Chervic may include Customer’s name/logo in client lists and marketing. Press releases or case studies require mutual approval.")
        
        add_heading("13.12 Export Compliance", 3)
        add_paragraph("Customer shall comply with applicable export control laws (U.S., India, others) and not export/re-export CAS Services without required licenses.")
        
        add_heading("13.13 No Third-Party Beneficiaries", 3)
        add_paragraph("No rights are granted to third parties, except permitted successors or assigns.")
        
        add_heading("13.14 Independent Contractors", 3)
        add_paragraph("Parties are independent contractors, not partners, joint venturers, or agents.")
        
        add_heading("13.15 Statistical Data Usage", 3)
        add_paragraph("Chervic may compile anonymized statistical data for internal improvements, ensuring no identification of Customer or users.")
        
        add_heading("13.16 Governing Law & Jurisdiction", 3)
        add_paragraph("Governed by Indian law, with exclusive jurisdiction in New Delhi courts.")
        
        add_heading("13.17 Regulatory & Legal Compliance", 3)
        add_paragraph("Both parties comply with applicable laws (data privacy, cybersecurity, taxation, labor). Chervic complies with India’s DPDP Act and cross-border data transfer regulations.")
        
        add_heading("13.18 Dispute Resolution", 3)
        add_paragraph("Disputes are resolved through mutual discussions within 15 days. If unresolved, escalate to arbitration or legal remedies under Indian law. Injunctive relief or IP issues may go to court.")
        
        add_heading("13.19 Counterparts & Electronic Signatures", 3)
        add_paragraph("This Agreement may be executed in counterparts (physical or electronic), each an original. E-signatures via email or platforms like DocuSign are valid.")
        
        add_heading("Signatures", 2)
        add_paragraph(f"**Chervic Advisory Services**\nSignature: __________________\nName: {content.get('chervic_name', '_______________________')}\nPosition: {content.get('chervic_title', '_______________________')}\nDate: {content.get('chervic_date', '_______________________')}")
        if content.get('chervic_signature') and os.path.exists(content['chervic_signature']):
            try:
                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(content['chervic_signature'], width=Inches(2))
                logger.info(f"Embedded Chervic signature in Signatures section: {content['chervic_signature']}")
            except Exception as e:
                logger.error(f"Error embedding Chervic signature in Signatures section: {e}")
                add_paragraph("Chervic Signature: [Error loading signature]")
        
        add_paragraph(f"\n**Customer**\nSignature: __________________\nName: {content.get('customer_sign_name', '_______________________')}\nPosition: {content.get('customer_sign_title', '_______________________')}\nDate: {content.get('customer_sign_date', '_______________________')}")
        if content.get('customer_signature') and os.path.exists(content['customer_signature']):
            try:
                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(content['customer_signature'], width=Inches(2))
                logger.info(f"Embedded Customer signature in Signatures section: {content['customer_signature']}")
            except Exception as e:
                logger.error(f"Error embedding Customer signature in Signatures section: {e}")
                add_paragraph("Customer Signature: [Error loading signature]")
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        logger.info("Successfully created document buffer")
        return buffer
    except Exception as e:
        logger.error(f"Error in create_document: {str(e)}")
        raise

def generate_pdf(content, pdf_filename):
    doc_buffer = create_document(content)
    docx_filename = f"temp_{uuid.uuid4()}.docx"
    docx_filepath = os.path.join(app.config['DOCX_DIR'], docx_filename).replace('\\', '/')
    pdf_filepath = os.path.join(app.config['OUTPUT_DIR'], pdf_filename).replace('\\', '/')
    
    try:
        with open(docx_filepath, 'wb') as f:
            f.write(doc_buffer.getvalue())
        pythoncom.CoInitialize()  # Add this line before convert
        try:
            convert(docx_filepath, pdf_filepath)
        finally:
            pythoncom.CoUninitialize()  # Add this line to clean up COM
        logger.info(f"Generated PDF: {pdf_filepath}")
    finally:
        if os.path.exists(docx_filepath):
            os.unlink(docx_filepath)
            logger.info(f"Deleted temporary docx: {docx_filepath}")
            
def create_table(doc, data, headers=None):
    rows = len(data) + (1 if headers else 0)
    cols = len(data[0]) if data else len(headers) if headers else 0
    if rows == 0 or cols == 0:
        logger.error("Invalid table dimensions")
        return
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    if headers:
        for j, header in enumerate(headers):
            if j < cols:  # Ensure valid column access
                table.rows[0].cells[j].text = header
    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            if j < cols:  # Ensure valid column access
                table.rows[i + (1 if headers else 0)].cells[j].text = str(cell_data)
            else:
                logger.warning(f"Invalid cell access at row {i + (1 if headers else 0)}, col {j}")
                

def extract_document_lines(doc_buffer):
    """Extract lines from a Word document buffer with page/section information."""
    try:
        doc = Document(doc_buffer)
        lines = []
        line_id = 0
        current_section = "Cover Page"  # Default to Cover Page
        section_map = {
            "Master Agreement": "Master Agreement",
            "Exhibit A: Plan & Pricing Schedule": "Exhibit A",
            "Exhibit B: Support & Maintenance Services": "Exhibit B",
            "Exhibit C: Service Level Agreement (SLA)": "Exhibit C",
            "CAS Agreement": "CAS Agreement"
        }

        # Process paragraphs directly from doc.paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                for line in para.text.split('\n'):
                    if line.strip():
                        # Check if paragraph is a heading to update section
                        for heading, section_name in section_map.items():
                            if line.strip().startswith(heading):
                                current_section = section_name
                                break
                        lines.append({
                            'id': f'line_{line_id}',
                            'text': line.strip(),
                            'type': 'paragraph',
                            'original': line.strip(),
                            'section': current_section
                        })
                        line_id += 1

        # Process tables directly from doc.tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            for line in para.text.split('\n'):
                                if line.strip():
                                    lines.append({
                                        'id': f'line_{line_id}',
                                        'text': line.strip(),
                                        'type': 'table_cell',
                                        'original': line.strip(),
                                        'section': current_section
                                    })
                                    line_id += 1

        logger.info(f"Extracted {len(lines)} lines from document")
        return lines
    except Exception as e:
        logger.error(f"Error extracting document lines: {e}")
        return []

def update_document_with_lines(doc_buffer, edited_lines):
    """Update the Word document with edited lines."""
    try:
        doc = Document(doc_buffer)
        line_id = 0
        for para in doc.paragraphs:
            if para.text.strip():
                new_text = []
                for line in para.text.split('\n'):
                    if line.strip():
                        line_key = f'line_{line_id}'
                        new_text.append(edited_lines.get(line_key, line.strip()))
                        line_id += 1
                    else:
                        new_text.append(line)
                para.text = '\n'.join(new_text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            new_text = []
                            for line in para.text.split('\n'):
                                if line.strip():
                                    line_key = f'line_{line_id}'
                                    new_text.append(edited_lines.get(line_key, line.strip()))
                                    line_id += 1
                                else:
                                    new_text.append(line)
                            para.text = '\n'.join(new_text)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        logger.info("Updated document with edited lines")
        return buffer
    except Exception as e:
        logger.error(f"Error updating document with lines: {e}")
        raise

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        if username in users and users[username] == password:
            session['username'] = username
            logger.info(f"User {username} logged in")
            return redirect(url_for('generate_nda'))
        else:
            flash("Invalid credentials", "error")
            logger.warning(f"Failed login attempt for username: {username}")
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.pop('username', None)
    logger.info("User logged out")
    flash("Logged out successfully", "success")
    return redirect(url_for('login'))

@app.route('/', methods=['GET', 'POST'])
def generate_nda():
    if 'username' not in session:
        return redirect(url_for('login'))
    
    # Default form data
    data = {
        "customer_name": "",
        "agreement_date": "",
        "start_date": "",
        "end_date": "",
        "billing_contact_name": "",
        "billing_email": "",
        "department": "",
        "billing_address": "",
        "tax_id": "",
        "po_number": "",
        "currency": "USD",
        "chervic_name": "",
        "chervic_title": "",
        "chervic_date": "",
        "customer_sign_name": "",
        "customer_sign_title": "",
        "customer_sign_date": "",
        "chervic_signature": "",
        "customer_signature": ""
    }
    
    # Load existing data
    if os.path.exists(app.config['JSON_FILE']):
        try:
            with open(app.config['JSON_FILE'], 'r') as f:
                data.update(json.load(f))
        except Exception as e:
            logger.error(f"Error loading JSON file: {e}")
            flash("Error loading saved data.", "error")
    
    if request.method == 'POST':
        # Validate required fields
        required_fields = ['customer_name', 'agreement_date', 'start_date', 'end_date', 
                         'billing_contact_name', 'billing_email', 'chervic_name', 
                         'chervic_title', 'chervic_date', 'customer_sign_name', 
                         'customer_sign_title', 'customer_sign_date']
        for field in required_fields:
            if not request.form.get(field):
                flash(f"{field.replace('_', ' ').title()} is required.", "error")
                return render_template('index.html', data=data)
        
        # Validate dates
        date_fields = ['agreement_date', 'start_date', 'end_date', 'chervic_date', 'customer_sign_date']
        for field in date_fields:
            if request.form.get(field) and not validate_date(request.form[field]):
                flash(f"Invalid date format for {field.replace('_', ' ').title()}. Use YYYY-MM-DD.", "error")
                return render_template('index.html', data=data)
        
        # Validate email
        if not re.match(r'^[\w\.-]+@[\w\.-]+\.\w+$', request.form.get('billing_email', '')):
            flash("Invalid billing email format.", "error")
            return render_template('index.html', data=data)
        
        # Handle signatures
        chervic_signature_path = data.get('chervic_signature', '')
        customer_signature_path = data.get('customer_signature', '')
        
        # Canvas signatures
        if request.form.get('chervic_signature_canvas'):
            chervic_signature_path = save_signature(request.form['chervic_signature_canvas'], 'chervic')
            if not chervic_signature_path:
                flash("Error saving Chervic signature.", "error")
                return render_template('index.html', data=data)
        
        if request.form.get('customer_signature_canvas'):
            customer_signature_path = save_signature(request.form['customer_signature_canvas'], 'customer')
            if not customer_signature_path:
                flash("Error saving Customer signature.", "error")
                return render_template('index.html', data=data)
        
        # Uploaded signatures
        if 'chervic_signature' in request.files and request.files['chervic_signature'].filename:
            file = request.files['chervic_signature']
            if file and allowed_file(file.filename):
                filename = f"chervic_{uuid.uuid4()}.{file.filename.rsplit('.', 1)[1].lower()}"
                chervic_signature_path = os.path.join(app.config['SIGNATURE_DIR'], filename).replace('\\', '/')
                file.save(chervic_signature_path)
                logger.info(f"Saved uploaded Chervic signature: {chervic_signature_path}")
            else:
                flash("Invalid Chervic signature file. Only PNG/JPEG allowed.", "error")
                return render_template('index.html', data=data)
        
        if 'customer_signature' in request.files and request.files['customer_signature'].filename:
            file = request.files['customer_signature']
            if file and allowed_file(file.filename):
                filename = f"customer_{uuid.uuid4()}.{file.filename.rsplit('.', 1)[1].lower()}"
                customer_signature_path = os.path.join(app.config['SIGNATURE_DIR'], filename).replace('\\', '/')
                file.save(customer_signature_path)
                logger.info(f"Saved uploaded Customer signature: {customer_signature_path}")
            else:
                flash("Invalid Customer signature file. Only PNG/JPEG allowed.", "error")
                return render_template('index.html', data=data)
        
        # Update data
        data = {
            "customer_name": sanitize_input(request.form['customer_name']),
            "agreement_date": request.form['agreement_date'],
            "start_date": request.form['start_date'],
            "end_date": request.form['end_date'],
            "billing_contact_name": sanitize_input(request.form['billing_contact_name']),
            "billing_email": request.form['billing_email'],
            "department": sanitize_input(request.form.get('department', '')),
            "billing_address": sanitize_input(request.form.get('billing_address', '')),
            "tax_id": sanitize_input(request.form.get('tax_id', '')),
            "po_number": sanitize_input(request.form.get('po_number', '')),
            "currency": request.form.get('currency', 'USD').upper() if request.form.get('currency', 'USD').upper() in ['INR', 'USD'] else 'USD',
            "chervic_name": sanitize_input(request.form['chervic_name']),
            "chervic_title": sanitize_input(request.form['chervic_title']),
            "chervic_date": request.form['chervic_date'],
            "customer_sign_name": sanitize_input(request.form['customer_sign_name']),
            "customer_sign_title": sanitize_input(request.form['customer_sign_title']),
            "customer_sign_date": request.form['customer_sign_date'],
            "chervic_signature": chervic_signature_path,
            "customer_signature": customer_signature_path
        }
        
        # Save to JSON
        try:
            with open(app.config['JSON_FILE'], 'w') as f:
                json.dump(data, f, indent=4)
            logger.info("Saved form data to JSON.")
        except Exception as e:
            logger.error(f"Error saving JSON file: {e}")
            flash("Error saving form data.", "error")
            return render_template('index.html', data=data)
        
        # Generate PDF
        pdf_filename = f"nda_{uuid.uuid4()}.pdf"
        try:
            data['pdf_filename'] = pdf_filename  # Add filename to content for line editing
            generate_pdf(data, pdf_filename)
            # Save metadata
            metadata = data.copy()
            metadata_file = os.path.join(app.config['METADATA_DIR'], f"{pdf_filename}_metadata.json").replace('\\', '/')
            with open(metadata_file, 'w') as f:
                json.dump(metadata, f, indent=4)
            # Save edit history
            save_edit_history(pdf_filename, session['username'], data)
            flash("PDF generated successfully.", "success")
            return redirect(url_for('view_pdf', filename=pdf_filename))
        except Exception as e:
            flash(f"Error generating PDF: {str(e)}", "error")
            return render_template('index.html', data=data)
    
    return render_template('index.html', data=data)

@app.route('/edit/<filename>', methods=['GET', 'POST'])
def edit_nda(filename):
    if 'username' not in session:
        return redirect(url_for('login'))
    
    metadata_file = os.path.join(app.config['METADATA_DIR'], f"{filename}_metadata.json").replace('\\', '/')
    if not os.path.exists(metadata_file):
        flash("Metadata not found.", "error")
        return redirect(url_for('generate_nda'))
    
    try:
        with open(metadata_file, 'r') as f:
            data = json.load(f)
        for key in ['chervic_signature', 'customer_signature']:
            if data.get(key):
                data[key] = data[key].replace('\\', '/')
    except Exception as e:
        logger.error(f"Error loading metadata: {e}")
        flash("Error loading metadata.", "error")
        return redirect(url_for('generate_nda'))
    
    history_file = os.path.join(app.config['EDIT_HISTORY_DIR'], f"{filename}_history.json").replace('\\', '/')
    edit_history = []
    if os.path.exists(history_file):
        try:
            with open(history_file, 'r') as f:
                edit_history = json.load(f)
        except Exception as e:
            logger.error(f"Error loading edit history: {e}")
            flash("Error loading edit history.", "error")
    
    if request.method == 'POST':
        # Validate required fields
        required_fields = ['customer_name', 'agreement_date', 'start_date', 'end_date', 
                         'billing_contact_name', 'billing_email', 'chervic_name', 
                         'chervic_title', 'chervic_date', 'customer_sign_name', 
                         'customer_sign_title', 'customer_sign_date']
        for field in required_fields:
            if not request.form.get(field):
                flash(f"{field.replace('_', ' ').title()} is required.", "error")
                return render_template('edit_nda.html', data=data, edit_history=edit_history)
        
        # Validate dates
        date_fields = ['agreement_date', 'start_date', 'end_date', 'chervic_date', 'customer_sign_date']
        for field in date_fields:
            if request.form.get(field) and not validate_date(request.form[field]):
                flash(f"Invalid date format for {field.replace('_', ' ').title()}. Use YYYY-MM-DD.", "error")
                return render_template('edit_nda.html', data=data, edit_history=edit_history)
        
        # Validate email
        if not re.match(r'^[\w\.-]+@[\w\.-]+\.\w+$', request.form.get('billing_email', '')):
            flash("Invalid billing email format.", "error")
            return render_template('edit_nda.html', data=data, edit_history=edit_history)
        
        # Handle signatures
        chervic_signature_path = data.get('chervic_signature', '')
        customer_signature_path = data.get('customer_signature', '')
        
        # Canvas signatures
        if request.form.get('chervic_signature_canvas'):
            chervic_signature_path = save_signature(request.form['chervic_signature_canvas'], 'chervic')
            if not chervic_signature_path:
                flash("Error saving Chervic signature.", "error")
                return render_template('edit_nda.html', data=data, edit_history=edit_history)
        
        if request.form.get('customer_signature_canvas'):
            customer_signature_path = save_signature(request.form['customer_signature_canvas'], 'customer')
            if not customer_signature_path:
                flash("Error saving Customer signature.", "error")
                return render_template('edit_nda.html', data=data, edit_history=edit_history)
        
        # Uploaded signatures
        if 'chervic_signature' in request.files and request.files['chervic_signature'].filename:
            file = request.files['chervic_signature']
            if file and allowed_file(file.filename):
                filename = f"chervic_{uuid.uuid4()}.{file.filename.rsplit('.', 1)[1].lower()}"
                chervic_signature_path = os.path.join(app.config['SIGNATURE_DIR'], filename).replace('\\', '/')
                file.save(chervic_signature_path)
                logger.info(f"Saved uploaded Chervic signature: {chervic_signature_path}")
            else:
                flash("Invalid Chervic signature file. Only PNG/JPEG allowed.", "error")
                return render_template('edit_nda.html', data=data, edit_history=edit_history)
        
        if 'customer_signature' in request.files and request.files['customer_signature'].filename:
            file = request.files['customer_signature']
            if file and allowed_file(file.filename):
                filename = f"customer_{uuid.uuid4()}.{file.filename.rsplit('.', 1)[1].lower()}"
                customer_signature_path = os.path.join(app.config['SIGNATURE_DIR'], filename).replace('\\', '/')
                file.save(customer_signature_path)
                logger.info(f"Saved uploaded Customer signature: {customer_signature_path}")
            else:
                flash("Invalid Customer signature file. Only PNG/JPEG allowed.", "error")
                return render_template('edit_nda.html', data=data, edit_history=edit_history)
        
        # Update data
        data = {
            "customer_name": sanitize_input(request.form['customer_name']),
            "agreement_date": request.form['agreement_date'],
            "start_date": request.form['start_date'],
            "end_date": request.form['end_date'],
            "billing_contact_name": sanitize_input(request.form['billing_contact_name']),
            "billing_email": request.form['billing_email'],
            "department": sanitize_input(request.form.get('department', '')),
            "billing_address": sanitize_input(request.form.get('billing_address', '')),
            "tax_id": sanitize_input(request.form.get('tax_id', '')),
            "po_number": sanitize_input(request.form.get('po_number', '')),
            "currency": request.form.get('currency', 'USD').upper() if request.form.get('currency', 'USD').upper() in ['INR', 'USD'] else 'USD',
            "chervic_name": sanitize_input(request.form['chervic_name']),
            "chervic_title": sanitize_input(request.form['chervic_title']),
            "chervic_date": request.form['chervic_date'],
            "customer_sign_name": sanitize_input(request.form['customer_sign_name']),
            "customer_sign_title": sanitize_input(request.form['customer_sign_title']),
            "customer_sign_date": request.form['customer_sign_date'],
            "chervic_signature": chervic_signature_path,
            "customer_signature": customer_signature_path,
            "pdf_filename": filename
        }
        
        # Save metadata
        try:
            with open(metadata_file, 'w') as f:
                json.dump(data, f, indent=4)
            logger.info(f"Updated metadata for {filename}")
        except Exception as e:
            logger.error(f"Error saving metadata: {e}")
            flash("Error saving metadata.", "error")
            return render_template('edit_nda.html', data=data, edit_history=edit_history)
        
        # Generate new PDF
        try:
            generate_pdf(data, filename)
            save_edit_history(filename, session['username'], data)
            flash("PDF updated successfully.", "success")
            return redirect(url_for('view_pdf', filename=filename))
        except Exception as e:
            flash(f"Error generating PDF: {str(e)}", "error")
            return render_template('edit_nda.html', data=data, edit_history=edit_history)
    
    return render_template('edit_nda.html', data=data, edit_history=edit_history)

@app.route('/pdf/<filename>')
def serve_pdf(filename):
    if 'username' not in session:
        flash("Please log in to view PDFs.", "error")
        return redirect(url_for('login'))
    filepath = os.path.normpath(os.path.join(app.config['OUTPUT_DIR'], filename))
    if os.path.exists(filepath):
        return send_file(filepath, mimetype='application/pdf')
    flash("PDF not found.", "error")
    logger.error(f"PDF not found: {filepath}")
    return redirect(url_for('generate_nda'))

@app.route('/view_pdf/<filename>')
def view_pdf(filename):
    if 'username' not in session:
        flash("Please log in to view PDFs.", "error")
        return redirect(url_for('login'))
    filepath = os.path.join(app.config['OUTPUT_DIR'], filename).replace('\\', '/')
    if os.path.exists(filepath):
        return render_template('view_pdf.html', pdf_filename=filename)
    flash("PDF not found.", "error")
    logger.error(f"PDF not found: {filepath}")
    return redirect(url_for('generate_nda'))

@app.route('/signature/<filename>')
def serve_signature(filename):
    if 'username' not in session:
        flash("Please log in to view signatures.", "error")
        return redirect(url_for('login'))
    filepath = os.path.join(app.config['SIGNATURE_DIR'], filename).replace('\\', '/')
    if os.path.exists(filepath):
        return send_file(filepath)
    flash("Signature not found.", "error")
    logger.error(f"Signature not found: {filepath}")
    return redirect(url_for('edit_nda', filename=request.args.get('filename', '')))

@app.route('/edit_lines/<filename>', methods=['GET', 'POST'])
def edit_nda_lines(filename):
    if 'username' not in session:
        flash("Please log in to edit lines.", "error")
        return redirect(url_for('login'))
    
    metadata_file = os.path.join(app.config['METADATA_DIR'], f"{filename}_metadata.json").replace('\\', '/')
    if not os.path.exists(metadata_file):
        flash("Metadata not found.", "error")
        return redirect(url_for('generate_nda'))
    
    try:
        with open(metadata_file, 'r') as f:
            metadata = json.load(f)
    except Exception as e:
        logger.error(f"Error loading metadata: {e}")
        flash("Error loading metadata.", "error")
        return redirect(url_for('generate_nda'))
    
    doc_buffer = create_document(metadata)
    lines = extract_document_lines(doc_buffer)
    
    lines_file = os.path.join(app.config['EDIT_HISTORY_DIR'], f"{filename}_lines.json").replace('\\', '/')
    edited_lines = {}
    if os.path.exists(lines_file):
        try:
            with open(lines_file, 'r') as f:
                edited_lines = json.load(f)
        except Exception as e:
            logger.error(f"Error loading line edits: {e}")
    
    if request.method == 'POST':
        action = request.form.get('action')
        if action == 'save_line':
            line_id = request.form.get('line_id')
            new_text = sanitize_input(request.form.get('text'))
            edited_lines[line_id] = new_text
            with open(lines_file, 'w') as f:
                json.dump(edited_lines, f, indent=4)
            logger.info(f"Saved edit for line {line_id}")
            flash("Line saved.", "success")
        elif action == 'cancel_line':
            line_id = request.form.get('line_id')
            if line_id in edited_lines:
                del edited_lines[line_id]
                with open(lines_file, 'w') as f:
                    json.dump(edited_lines, f, indent=4)
                logger.info(f"Cancelled edit for line {line_id}")
                flash("Line edit cancelled.", "success")
        elif action == 'save_all':
            updated_doc_buffer = update_document_with_lines(doc_buffer, edited_lines)
            docx_filename = f"temp_{uuid.uuid4()}.docx"
            docx_filepath = os.path.join(app.config['DOCX_DIR'], docx_filename).replace('\\', '/')
            pdf_filepath = os.path.join(app.config['OUTPUT_DIR'], filename).replace('\\', '/')
            try:
                with open(docx_filepath, 'wb') as f:
                    f.write(updated_doc_buffer.getvalue())
                pythoncom.CoInitialize()  # Add this line
                try:
                    convert(docx_filepath, pdf_filepath)  # Line 1157
                finally:
                    pythoncom.CoUninitialize()  # Add this line
                os.unlink(docx_filepath)
                save_edit_history(filename, session['username'], {'lines': edited_lines})
                logger.info(f"Regenerated PDF with all line edits: {pdf_filepath}")
                flash("All changes saved and PDF regenerated.", "success")
                return redirect(url_for('view_pdf', filename=filename))
            except Exception as e:
                logger.error(f"Error regenerating PDF: {e}")
                flash(f"Error regenerating PDF: {str(e)}", "error")
        
        return redirect(url_for('edit_nda_lines', filename=filename))
    
    return render_template('edit_lines.html', filename=filename, lines=lines, edited_lines=edited_lines)

if __name__ == '__main__':
    app.run(debug=True)
